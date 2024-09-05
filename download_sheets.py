from google.oauth2 import service_account
from googleapiclient.discovery import build
import time
import random
from googleapiclient.errors import HttpError
import httplib2
import re
import io
from googleapiclient.http import MediaIoBaseDownload
from docx import Document
import os



#change this to the location of your credentials1.json file 
SERVICE_ACCOUNT_FILE = '/Users/g24isha/FormGeneration/credentials1.json'

# Scopes required for Google Drive and Forms API
SCOPES = [
    'https://www.googleapis.com/auth/drive',
    'https://www.googleapis.com/auth/forms.body'
]

# Authenticate and create the service
credentials = service_account.Credentials.from_service_account_file(
    SERVICE_ACCOUNT_FILE, scopes=SCOPES)
http = httplib2.Http(timeout=120)
drive_service = build('drive', 'v3', credentials=credentials)
forms_service = build('forms', 'v1', credentials=credentials)
sheets_service = build('sheets', 'v4', credentials=credentials)

MAX_RETRIES = 8
INITIAL_DELAY = 1
MAX_DELAY = 100

def exponential_backoff(retry_count):
    return min(MAX_DELAY, INITIAL_DELAY * (2 ** retry_count) + random.uniform(0, 1))

def execute_with_backoff(func, *args, **kwargs):
    retry_count = 0
    while retry_count < MAX_RETRIES:
        try:
            result = func(*args, **kwargs)
            if callable(getattr(result, 'execute', None)):
                result = result.execute()
            return result
        except HttpError as e:
            if 'Quota exceeded' in str(e) or 'User rate limit exceeded' in str(e):
                retry_count += 1
                delay = exponential_backoff(retry_count)
                print(f"Quota exceeded, retrying in {delay:.2f} seconds... (attempt {retry_count}/{MAX_RETRIES})")
                time.sleep(delay)
            else:
                print(f"HTTP Exception: {e}")
                raise e
        except TimeoutError as e:
            retry_count += 1
            delay = exponential_backoff(retry_count)
            print(f"Timeout error, retrying in {delay:.2f} seconds... (attempt {retry_count}/{MAX_RETRIES})")
            time.sleep(delay)
        except Exception as e:
            print(f"Exception: {e}")
            raise e
    print(f"Failed after {MAX_RETRIES} retries.")
    return None

def download_file_with_backoff(file_id):
    request = drive_service.files().get_media(fileId=file_id)
    fh = io.BytesIO()
    downloader = MediaIoBaseDownload(fh, request)
    done = False
    while not done:
        status, done = downloader.next_chunk()
    return fh.getvalue()

def get_topic(doc_id):
    try:
        file_content = execute_with_backoff(download_file_with_backoff, doc_id)
    except Exception as e:
        print(f"Failed to download file {file_name}: {e}")
        return
    # Save the downloaded file temporarily
    local_file_path = f"{file_name}.docx"
    with open(local_file_path, 'wb') as f:
        f.write(file_content)
        
    doc_content = Document(local_file_path)
    extracted_text = "\n".join([paragraph.text for paragraph in doc_content.paragraphs])
    extracted_text = re.sub(r'\s+', ' ', extracted_text).strip()
    

    topic_beg = "Please indicate your choice of topic:"
    topic_end  = "Before we start writing the actual idea, please also indicate how familiar you are with the given topic on a scale of 1 - 5 (this is just for us to understand potential confounders):"
    index1 = extracted_text.find(topic_beg) + 37
    index2 = extracted_text.find(topic_end)
    topic = extracted_text[index1:index2+1]
    topic = re.sub(r'\D', '', topic)
   
    file = drive_service.files().get(fileId=file_id, fields='webViewLink').execute()
    if len(topic) < 1:
        print("NO TOPIC FOR: ", file_name)
        return file['webViewLink'], "[No topic was returned]"
    topic_num = int(topic[0])
    os.remove(local_file_path)

#create lists of scores, topics, and reviewers
SPREADSHEET_ID_3 = ''
TOPICS_RANGE = 'AllIdeas!A:A'
CONDITIONS_RANGE = 'AllIdeas!B:B'
RESPONSES_RANGE = 'AllIdeas!I:I'

result = execute_with_backoff(sheets_service.spreadsheets().values().get, spreadsheetId=SPREADSHEET_ID, range=TOPICS_RANGE)
result2 = execute_with_backoff(sheets_service.spreadsheets().values().get, spreadsheetId=SPREADSHEET_ID, range=CONDITIONS_RANGE)
result3 = execute_with_backoff(sheets_service.spreadsheets().values().get, spreadsheetId=SPREADSHEET_ID, range=RESPONSES_RANGE)

topic_results = result.get('values', [])
conditions_results = result2.get('values', [])
responses_results = result3.get('values', [])
file_num = 1
for i in range(1, len(topic_results)):
    if len(responses_results[i]) == 0:
        continue
    topic = topic_results[i][0]
    hashtag = topic.find('#')
    topic = topic[:hashtag-1]
    
    author = conditions_results[i][0]
    if author.strip().lower() == 'ai + ai':
        condition = 'AI + AI'
    elif author.find('AI + Human') != -1:
        condition = 'AI + Human'
    else:
        condition = 'Human'
        
    id = responses_results[i][0]
    first_index = id.find('/d/')
    second_index = id.find('/edit')
    id = id[first_index+3:second_index]
    try:
        file_content = execute_with_backoff(download_file_with_backoff, id)
    except Exception as e:
        print(f"Failed to download file with id {id}: {e}")
        
    # Save the downloaded file temporarily
    local_file_path = f"{file_num}.csv"
    # with open(local_file_path, 'wb') as f:
    #     f.write(file_content)
        
    # doc_content = Document(local_file_path)
    # extracted_text = "\n".join([paragraph.text for paragraph in doc_content.paragraphs])
    # extracted_text = re.sub(r'\s+', ' ', extracted_text).strip()
    

    # topic_beg = "Please indicate your choice of topic:"
    # topic_end  = "Before we start writing the actual idea, please also indicate how familiar you are with the given topic on a scale of 1 - 5 (this is just for us to understand potential confounders):"
    # index1 = extracted_text.find(topic_beg) + 37
    # index2 = extracted_text.find(topic_end)
    # topic = extracted_text[index1:index2+1]
    # topic = re.sub(r'\D', '', topic)
   
    # file = drive_service.files().get(fileId=file_id, fields='webViewLink').execute()
    # if len(topic) < 1:
    #     print("NO TOPIC FOR: ", file_name)
    #     return file['webViewLink'], "[No topic was returned]"
    # topic_num = int(topic[0])
    os.remove(local_file_path)
    file_num += 1
    # url = f"https://docs.google.com/spreadsheets/d/{id}/export?format=csv"
    # response = requests.get(url)
    
    # filename = f"{sheet_id}.csv"
    # with open(filename, 'wb') as file:
    #     file.write(response.content)
    # print(f"Downloaded the sheet as {filename}")